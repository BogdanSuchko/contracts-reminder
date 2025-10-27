from __future__ import annotations

from pathlib import Path

from docx import Document

TEMPLATES = Path("templates")
TEMPLATES.mkdir(parents=True, exist_ok=True)


def create_termination_template() -> None:
    doc = Document()
    doc.add_paragraph("{{ organization }}").alignment = 0
    doc.add_paragraph("УВЕДОМЛЕНИЕ").alignment = 1
    doc.add_paragraph("№ {{ document_number }}").alignment = 1
    doc.add_paragraph(" ")
    doc.add_paragraph(
        "В связи с истечением срока действия контракта от {{ contract_date }} № {{ contract_number }} "
        "уведомляю Вас о намерении прекратить трудовые отношения с {{ employee_full_name }} с {{ contract_end_date }}."
    )
    doc.add_paragraph(" ")
    doc.add_paragraph("Директор ____________________ {{ director_signature }}")
    doc.add_paragraph("{{ director_name }}")
    doc.add_paragraph(" ")
    doc.add_paragraph("Визы")
    doc.save(TEMPLATES / "notify_termination.docx")


def create_extension_template() -> None:
    doc = Document()
    doc.add_paragraph("{{ organization }}")
    doc.add_paragraph("УВЕДОМЛЕНИЕ № {{ document_number }}").alignment = 1
    doc.add_paragraph("г. Минск").alignment = 1
    doc.add_paragraph(" ")
    doc.add_paragraph("{{ position }}")
    doc.add_paragraph("{{ employee_full_name }}")
    doc.add_paragraph(" ")
    doc.add_paragraph("О продлении трудового договора (контракта)").alignment = 1
    doc.add_paragraph(" ")
    doc.add_paragraph(
        "Срок трудового договора (контракта) от {{ contract_date }} № {{ contract_number }}, заключенного "
        "между {{ organization }} и Вами."
    )
    doc.add_paragraph(
        "Срок действующего трудового договора (контракта) истекает {{ contract_end_date }}."
    )
    doc.add_paragraph(
        "Предлагаем Вам продлить действующий трудовой договор (контракт) на {{ extension_term }} "
        "с {{ extension_start }} по {{ extension_end }}."
    )
    doc.add_paragraph(
        "Просим сообщить о согласии или несогласии на продление до {{ response_deadline }}."
    )
    doc.add_paragraph(" ")
    doc.add_paragraph("Директор ____________________ {{ director_signature }}")
    doc.add_paragraph("{{ director_name }}")
    doc.save(TEMPLATES / "notify_extension.docx")


if __name__ == "__main__":
    create_termination_template()
    create_extension_template()
